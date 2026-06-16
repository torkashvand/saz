"""Derive a tokenized RFQ template (and placeholder map) from the original
GÉANT RFQ .docx. Highlighted spans become {{token}} markers; original
formatting is preserved by reusing the first run of each span."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


def _is_highlighted(run) -> bool:
    color = run.font.highlight_color
    return color is not None and str(color).lower() != "none"


def _match_token(text: str, hints: list[tuple[str, str]]) -> str | None:
    low = text.lower()
    for substring, token in hints:
        if substring.lower() in low:
            return token
    return None


def _replace_in_paragraph(
    paragraph: Paragraph,
    hints: list[tuple[str, str]],
    counters: dict[str, int],
    label_unmatched: bool,
) -> int:
    """Collapse each contiguous highlighted run-span into one {{token}} run.

    A span matching a hint becomes ``{{token}}``. An unmatched span becomes
    ``{{extra_N}}`` when ``label_unmatched`` is set (inspection mode), otherwise
    it is left untouched so the template keeps its original example text.
    Returns the number of spans turned into tokens.
    """
    runs = paragraph.runs
    replaced = 0
    i = 0
    while i < len(runs):
        if not _is_highlighted(runs[i]):
            i += 1
            continue
        j = i
        while j + 1 < len(runs) and _is_highlighted(runs[j + 1]):
            j += 1
        span_text = "".join(r.text for r in runs[i : j + 1])
        token = _match_token(span_text, hints)
        if token is None:
            if not label_unmatched:
                i = j + 1
                continue
            counters["extra"] += 1
            token = f"extra_{counters['extra']}"
        runs[i].text = f"{{{{{token}}}}}"
        for k in range(i + 1, j + 1):
            runs[k].text = ""
        replaced += 1
        i = j + 1
    return replaced


def replace_highlighted_spans(
    doc, hints: list[tuple[str, str]], label_unmatched: bool = True
) -> int:
    """Replace highlighted spans in the document body with tokens.

    When ``label_unmatched`` is True (default), spans with no matching hint
    become ``{{extra_N}}`` markers, useful for discovering what the template
    contains. When False, only hint-matched spans are tokenized and all other
    highlighted text is left as-is. Returns the number of spans tokenized.
    """
    counters = {"extra": 0}
    total = 0
    for paragraph in doc.paragraphs:
        total += _replace_in_paragraph(paragraph, hints, counters, label_unmatched)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += _replace_in_paragraph(paragraph, hints, counters, label_unmatched)
    return total


def build(source: Path, out_docx: Path, out_map: Path) -> dict[str, int]:
    from saz.examples.templates.rfq_tokens import (
        RFQ_TOKENS,
        TOKEN_NOTES,
        TOKEN_SOURCE_HINTS,
    )

    doc = Document(str(source))
    replaced = replace_highlighted_spans(doc, TOKEN_SOURCE_HINTS, label_unmatched=False)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))

    # Which canonical tokens actually landed in the template?
    rendered = Document(str(out_docx))
    body_text = "\n".join(p.text for p in rendered.paragraphs)
    for table in rendered.tables:
        for row in table.rows:
            for cell in row.cells:
                body_text += "\n" + cell.text
    present = [t for t in RFQ_TOKENS if f"{{{{{t}}}}}" in body_text]
    missing = [t for t in RFQ_TOKENS if t not in present]

    lines = ["# RFQ Placeholder Map", "", "| Token | In template | Notes |", "|---|---|---|"]
    for t in RFQ_TOKENS:
        note = TOKEN_NOTES.get(t, "")
        lines.append(f"| `{{{{{t}}}}}` | {'yes' if t in present else 'NO'} | {note} |")
    out_map.parent.mkdir(parents=True, exist_ok=True)
    out_map.write_text("\n".join(lines) + "\n", encoding="utf-8")

    return {"replaced": replaced, "present": len(present), "missing": len(missing)}


def main() -> None:
    parser = argparse.ArgumentParser(description="Build tokenized RFQ template.")
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument(
        "--out-docx",
        type=Path,
        default=Path(__file__).parent / "rfq_template.docx",
    )
    parser.add_argument(
        "--out-map",
        type=Path,
        default=Path(__file__).resolve().parents[4]
        / "docs"
        / "procurement"
        / "rfq_placeholder_map.md",
    )
    args = parser.parse_args()
    stats = build(args.source, args.out_docx, args.out_map)
    print(f"replaced={stats['replaced']} present={stats['present']} missing={stats['missing']}")


if __name__ == "__main__":
    main()
