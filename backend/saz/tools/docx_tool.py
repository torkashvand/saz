"""Document render tool: fill {{token}} markers in a .docx template."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any
from uuid import uuid4

import structlog
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

logger = structlog.get_logger(__name__)

_TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


class DocxRenderTool:
    """Fill {{token}} markers in a .docx and write the result to storage."""

    def __init__(self, storage_path: str = "/tmp/saz/artifacts"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.logger = logger.bind(tool="docx_render")

    @property
    def spec(self) -> dict[str, Any]:
        return {
            "name": "docx_render",
            "description": "Fill {{token}} placeholders in a .docx template and store the result.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "template": {"type": "string", "description": "Path to the tokenized .docx."},
                    "values": {
                        "type": "object",
                        "description": "Map of token name -> replacement string.",
                        "additionalProperties": True,
                    },
                    "output_name": {"type": "string", "description": "Name for the rendered file."},
                    "require_all": {
                        "type": "boolean",
                        "default": True,
                        "description": "Fail if any token is left unfilled.",
                    },
                },
                "required": ["template", "values", "output_name"],
            },
        }

    def _fill_paragraph(self, paragraph: Paragraph, values: dict[str, str]) -> int:
        filled = 0
        for run in paragraph.runs:
            if "{{" not in run.text:
                continue

            def repl(match: re.Match[str]) -> str:
                nonlocal filled
                token = match.group(1)
                if token in values:
                    filled += 1
                    return str(values[token])
                return match.group(0)

            run.text = _TOKEN_RE.sub(repl, run.text)
        return filled

    @staticmethod
    def _remaining_tokens(doc: DocxDocument) -> list[str]:
        found: list[str] = []

        def scan(paragraph: Paragraph) -> None:
            for m in _TOKEN_RE.finditer(paragraph.text):
                if m.group(1) not in found:
                    found.append(m.group(1))

        for p in doc.paragraphs:
            scan(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        scan(p)
        return found

    async def render(
        self,
        template: str,
        values: dict[str, Any],
        output_name: str,
        require_all: bool = True,
        run_id: str = "",
        step_id: str = "",
    ) -> dict[str, Any]:
        template_path = Path(template)
        if not template_path.is_absolute() and not template_path.exists():
            # Bundled templates are referenced relative to the backend root
            # (e.g. "saz/examples/templates/rfq_template.docx"); resolve against
            # the package root so rendering does not depend on the process CWD.
            import saz

            backend_root = Path(saz.__file__).resolve().parent.parent
            candidate = backend_root / template
            if candidate.exists():
                template_path = candidate
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template}")

        str_values = {k: ("" if v is None else str(v)) for k, v in values.items()}
        doc = Document(str(template_path))

        # Tokens present in the template before substitution, and which of them
        # received an empty/whitespace (or missing) value. An empty fill of a
        # mandatory token is a failure — never silently emit a blank section.
        present_tokens = self._remaining_tokens(doc)
        empty_tokens = [t for t in present_tokens if not str_values.get(t, "").strip()]

        filled = 0
        for p in doc.paragraphs:
            filled += self._fill_paragraph(p, str_values)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        filled += self._fill_paragraph(p, str_values)

        # Unfilled = literal {{token}} still in the doc (missing value, or a token
        # split across runs) PLUS tokens that were filled with an empty value.
        unfilled = sorted(set(self._remaining_tokens(doc)) | set(empty_tokens))
        if require_all and unfilled:
            raise ValueError(f"Unfilled or empty mandatory tokens: {', '.join(unfilled)}")

        artifact_id = str(uuid4())
        safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", output_name) or "document"
        out_path = self.storage_path / f"{artifact_id}_{safe_name}.docx"
        doc.save(str(out_path))

        self.logger.info(
            "docx.rendered",
            artifact_id=artifact_id,
            output_name=output_name,
            filled=filled,
            unfilled=len(unfilled),
            run_id=run_id,
            step_id=step_id,
            path=str(out_path),
        )

        return {
            "artifact_id": artifact_id,
            "name": output_name,
            "path": str(out_path),
            "filled": filled,
            "unfilled": unfilled,
            "byte_size": out_path.stat().st_size,
        }
