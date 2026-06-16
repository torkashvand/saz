from pathlib import Path

import pytest
from docx import Document

from saz.tools.docx_tool import DocxRenderTool

_TEMPLATE = (
    Path(__file__).resolve().parents[2] / "saz" / "examples" / "templates" / "rfq_template.docx"
)

_VALUES = {
    "title_system_name": "HR Information System",
    "date_of_issue": "05/07/2024",
    "version": "1.0",
    "reference_number": "T88815",
    "contact_name": "Badreddine Ajbar El Gueriri",
    "contact_role": "Buyer",
    "contact_phone": "+31 6 29003633",
    "contact_email": "badre.ajbar@geant.org",
    "background": "GÉANT seeks to replace its existing HR Information System.",
    "objective": "Find a cost-effective, modern, well-supported HR system.",
    "scope": "A core HR system with optional modules for future expansion.",
    "plan_rfq_issued": "05/07/2024",
    "plan_clarification_deadline": "15/07/2024",
    "plan_response_deadline": "19/07/2024",
    "plan_eval1_end": "02/08/2024",
    "plan_eval2_end": "09/08/2024",
    "plan_awarding": "15/08/2024",
    "plan_commencement": "01/10/2024",
    "minimum_requirements": "1. SSO via OIDC/SAML2\n2. EU/UK data residency\n3. TLS 1.2+",
    "weight_qualitative": "80%",
    "q1_weight": "50%",
    "q2_weight": "20%",
    "q3_weight": "10%",
    "weight_price": "20%",
    "budget_cap_licenses": "20.000",
    "budget_cap_implementation": "10.000",
}


@pytest.mark.asyncio
async def test_real_template_renders_with_all_present_tokens(tmp_path):
    tool = DocxRenderTool(storage_path=str(tmp_path))
    result = await tool.render(
        template=str(_TEMPLATE),
        values=_VALUES,
        output_name="rfq_final_T88815",
        require_all=True,
    )
    assert result["unfilled"] == []
    doc = Document(result["path"])

    full = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full += "\n" + cell.text

    assert "T88815" in full  # reference number rendered
    assert "Find a cost-effective" in full  # objective narrative rendered
    assert "{{" not in full  # no leftover token markers anywhere
