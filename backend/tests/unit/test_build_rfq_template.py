from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from saz.examples.templates.build_rfq_template import replace_highlighted_spans


def _highlight(run):
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def test_collapses_contiguous_highlighted_runs_into_one_token():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Objective: ")  # not highlighted
    r1 = p.add_run("find a HR system ")
    r2 = p.add_run("that is cost effective")
    _highlight(r1)
    _highlight(r2)

    n = replace_highlighted_spans(doc, hints=[("find a hr system", "objective")])

    assert n == 1
    text = "".join(r.text for r in p.runs)
    assert text == "Objective: {{objective}}"


def test_unmatched_span_gets_extra_token_and_is_reported():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Some variable value")
    _highlight(r)

    n = replace_highlighted_spans(doc, hints=[])

    assert n == 1
    assert "{{extra_1}}" in "".join(run.text for run in p.runs)


def test_two_unmatched_spans_increment_extra_counter():
    doc = Document()
    p1 = doc.add_paragraph()
    _highlight(p1.add_run("first variable"))
    p2 = doc.add_paragraph()
    _highlight(p2.add_run("second variable"))

    n = replace_highlighted_spans(doc, hints=[])

    assert n == 2
    assert "{{extra_1}}" in "".join(r.text for r in p1.runs)
    assert "{{extra_2}}" in "".join(r.text for r in p2.runs)


def test_first_matching_hint_wins():
    doc = Document()
    p = doc.add_paragraph()
    _highlight(p.add_run("RFQ issued: 05/07/2024"))

    # Both hints' substrings are present; the earlier hint must win.
    n = replace_highlighted_spans(
        doc,
        hints=[("rfq issued", "plan_rfq_issued"), ("05/07/2024", "date_of_issue")],
    )

    assert n == 1
    assert "{{plan_rfq_issued}}" in "".join(r.text for r in p.runs)


def test_tokenizes_highlighted_span_inside_table_cell():
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    _highlight(cell.paragraphs[0].add_run("T88815"))

    n = replace_highlighted_spans(doc, hints=[("t88815", "reference_number")])

    assert n == 1
    assert "{{reference_number}}" in cell.text
