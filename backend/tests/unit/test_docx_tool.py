from pathlib import Path

import pytest
from docx import Document

from saz.tools.docx_tool import DocxRenderTool


def _make_template(tmp_path, body: str):
    path = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_paragraph(body)
    doc.save(str(path))
    return str(path)


@pytest.mark.asyncio
async def test_fills_tokens_and_stores_artifact(tmp_path):
    tpl = _make_template(tmp_path, "Objective: {{objective}}")
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))

    result = await tool.render(
        template=tpl,
        values={"objective": "Find a modern HR system"},
        output_name="draft_rfq",
        run_id="r1",
        step_id="s1",
    )

    assert result["filled"] == 1
    assert result["unfilled"] == []
    rendered = Document(result["path"])
    assert rendered.paragraphs[0].text == "Objective: Find a modern HR system"


@pytest.mark.asyncio
async def test_reports_unfilled_and_fails_when_require_all(tmp_path):
    tpl = _make_template(tmp_path, "A {{objective}} B {{scope}}")
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))

    with pytest.raises(ValueError) as exc:
        await tool.render(
            template=tpl,
            values={"objective": "x"},
            output_name="draft",
            require_all=True,
        )
    assert "scope" in str(exc.value)


@pytest.mark.asyncio
async def test_reports_unfilled_without_failing_when_not_require_all(tmp_path):
    tpl = _make_template(tmp_path, "A {{objective}} B {{scope}}")
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))

    result = await tool.render(
        template=tpl, values={"objective": "x"}, output_name="draft", require_all=False
    )
    assert result["unfilled"] == ["scope"]


@pytest.mark.asyncio
async def test_missing_template_raises(tmp_path):
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))
    with pytest.raises(FileNotFoundError):
        await tool.render(template=str(tmp_path / "nope.docx"), values={}, output_name="x")


@pytest.mark.asyncio
async def test_fills_tokens_inside_table_cells(tmp_path):
    path = tmp_path / "tbl.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run("Ref: {{reference_number}}")
    table.cell(0, 1).paragraphs[0].add_run("Ver: {{version}}")
    doc.save(str(path))
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))

    result = await tool.render(
        template=str(path),
        values={"reference_number": "T88815"},
        output_name="tbl",
        require_all=False,
    )

    assert result["unfilled"] == ["version"]  # unfilled token in a cell is reported
    rendered = Document(result["path"])
    assert rendered.tables[0].cell(0, 0).text == "Ref: T88815"


@pytest.mark.asyncio
async def test_token_split_across_runs_is_not_filled_but_is_detected(tmp_path):
    # The fill is per-run; detection is per-paragraph. A token split across runs
    # must NOT silently pass require_all.
    path = tmp_path / "split.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{")
    p.add_run("objective}}")
    doc.save(str(path))
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))

    with pytest.raises(ValueError) as exc:
        await tool.render(
            template=str(path),
            values={"objective": "x"},
            output_name="split",
            require_all=True,
        )
    assert "objective" in str(exc.value)

    result = await tool.render(
        template=str(path),
        values={"objective": "x"},
        output_name="split",
        require_all=False,
    )
    assert result["unfilled"] == ["objective"]


@pytest.mark.asyncio
async def test_output_name_path_separators_are_sanitized(tmp_path):
    tpl = _make_template(tmp_path, "Hi {{x}}")
    tool = DocxRenderTool(storage_path=str(tmp_path / "art"))

    result = await tool.render(template=tpl, values={"x": "y"}, output_name="../escape/evil")

    # Separators are stripped so the file lands directly in storage_path
    # (no traversal escape), regardless of the requested name.
    out = Path(result["path"])
    assert out.parent == tool.storage_path
    assert "/" not in out.name
    assert out.name.endswith(".docx")


def test_docx_render_in_default_registry():
    from saz.tools.registry import create_default_registry

    registry = create_default_registry(enable_ai_ops=False)
    assert "docx_render" in registry.list_tools()
    assert registry.get_tool_spec("docx_render")["name"] == "docx_render"
